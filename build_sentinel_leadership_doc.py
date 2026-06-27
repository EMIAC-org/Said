from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


SKILL_SCRIPTS = Path(
    "/Users/anishsuman/.codex/plugins/cache/openai-primary-runtime/documents/26.430.10722/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


OUT = Path("/Users/anishsuman/Documents/projects/emiac/said/Sentinel_Stryker_Leadership_Pitch.docx")

ACCENT = RGBColor(30, 64, 175)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
LIGHT = "F3F6FB"
LIGHT_BLUE = "EFF6FF"
RULE = "D9DEE8"
GREEN = RGBColor(20, 83, 45)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=RULE, size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_border(paragraph, color=RULE):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_end)


def add_text(paragraph, text, bold=False, color=None, size=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_heading(doc, text, level=1, page_break_before=False):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.page_break_before = page_break_before
    add_text(paragraph, text)
    return paragraph


def add_body(doc, text, after=6):
    paragraph = doc.add_paragraph()
    paragraph.style = "Normal"
    paragraph.paragraph_format.space_after = Pt(after)
    add_text(paragraph, text)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Twips(720)
        paragraph.paragraph_format.first_line_indent = Twips(-360)
        paragraph.paragraph_format.space_after = Pt(7)
        add_text(paragraph, item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Twips(720)
        paragraph.paragraph_format.first_line_indent = Twips(-360)
        paragraph.paragraph_format.space_after = Pt(7)
        add_text(paragraph, item)


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "BFD7FF")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, title, bold=True, color=ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(p2, body, color=INK)
    apply_table_geometry(table, [9360], table_width_dxa=9360, indent_dxa=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, weights, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, text, bold=True, color=INK, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_border(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, text, color=INK, size=font_size)
    widths = column_widths_from_weights(weights, 9360)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title.font.size = Pt(23)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)

    for level, size in [(1, 15), (2, 12.5), (3, 11)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT if level == 1 else INK
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(5)

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(7)

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(2)
    add_text(hp, "Sentinel leadership pitch | Research-backed memo", color=MUTED, size=8.5)
    set_paragraph_border(hp)
    footer = section.footer
    add_page_number(footer.paragraphs[0])
    return doc


def build():
    doc = setup_document()

    p = doc.add_paragraph(style="Title")
    add_text(p, "Sentinel: AI Quality Gate for Regulated Software Delivery")
    sp = doc.add_paragraph(style="Subtitle")
    add_text(
        sp,
        "A simple, honest, research-backed pitch for Stryker SGTC leadership | Prepared May 2026",
    )

    meta = [
        ("Audience", "Stryker SGTC and product engineering leadership"),
        ("Purpose", "Explain why Sentinel is different, credible, and worth a focused pilot"),
        ("Tone", "Positive, practical, and transparent about risks"),
    ]
    add_table(doc, ["Field", "Summary"], meta, [1.2, 4.8], 9.5)

    add_callout(
        doc,
        "Recommended leadership message",
        "Sentinel is not another AI coding assistant. It is a private, reviewable quality gate that helps Stryker move faster while strengthening the evidence trail expected in regulated software work.",
    )

    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "The market already has strong AI tools for code review, unit-test generation, visual testing, and QA automation. That is exactly why Sentinel should not be pitched as a generic AI test generator. The more credible position is narrower and stronger: Sentinel verifies every pull request inside Stryker's own environment, proposes only tests that survive strict quality gates, and leaves behind audit-ready evidence.",
    )
    add_body(
        doc,
        "This message should feel authentic to leadership because it respects the reality of the competitor landscape. Diffblue, Qodo, CodeRabbit, mabl, Tricentis, Launchable, Applitools, and GitHub Copilot all solve important parts of the quality problem. Sentinel stands out by combining PR-native test creation, gated verification, private Azure deployment, medical software traceability, and human review.",
    )

    add_heading(doc, "The One-Sentence Pitch", 1)
    add_callout(
        doc,
        "Pitch line",
        "Sentinel helps Stryker ship software faster without weakening review discipline, test evidence, or regulated delivery confidence.",
        "F0FDF4",
    )

    add_heading(doc, "Why This Matters To Stryker", 1)
    add_bullets(
        doc,
        [
            "Stryker is expanding digital and connected-care offerings, including the 2026 SmartHospital Platform, where software quality and workflow trust are strategic assets.",
            "SGTC Gurugram is positioned by Stryker as a global R&D engine across divisions, which makes it a strong home for a platform capability that can scale across teams.",
            "AI coding tools are increasing development speed. Leadership will naturally ask whether quality controls and evidence creation are improving at the same pace.",
            "Medical-device software work needs more than speed. It needs traceability, reviewability, controlled change, and confidence that automation can be explained later.",
        ],
    )

    add_heading(doc, "What Competitors Already Do Well", 1)
    add_body(
        doc,
        "This is the honest market picture. The goal is not to pretend competitors are weak. The goal is to show that Sentinel is aimed at a different leadership concern.",
    )
    add_heading(doc, "Competitor Reality", 1, page_break_before=True)
    competitors = [
        (
            "Diffblue Cover",
            "Autonomous Java/Kotlin unit tests in IDE, CLI, and CI; can run locally or in CI.",
            "Very strong, but narrower platform scope. Sentinel should win on multi-platform PR assurance and regulated evidence.",
        ),
        (
            "Qodo",
            "AI PR review, rules, governance, and PR-based test generation across many languages.",
            "Strong review platform. Sentinel should emphasize verified test artifacts and in-perimeter quality gates, not just comments.",
        ),
        (
            "CodeRabbit",
            "Context-aware AI code review, one-click fixes, and beta unit-test generation.",
            "Polished developer workflow. Sentinel should win with Azure-first private deployment, audit trail, and medical SDLC framing.",
        ),
        (
            "GitHub Copilot",
            "Default enterprise coding agent inside GitHub; can open or modify PRs.",
            "Useful coding agent, but generic. Sentinel is a controlled quality gate, not a code producer.",
        ),
        (
            "mabl / Tricentis",
            "Mature AI and enterprise QA automation for UI, API, E2E, auto-healing, and risk-based testing.",
            "Strong QA layer. Sentinel is earlier in the flow: developer PR quality before downstream regression cycles.",
        ),
        (
            "Launchable / Applitools",
            "Predictive test selection and visual AI regression testing.",
            "Excellent specialist tools. Sentinel can integrate similar signals but owns the PR-level evidence workflow.",
        ),
    ]
    add_table(doc, ["Competitor", "Strength", "Sentinel wedge"], competitors, [1.25, 2.45, 2.7], 8.3)

    add_heading(doc, "What Makes Sentinel Stand Out", 1, page_break_before=True)
    add_numbered(
        doc,
        [
            "Private by design: source context stays inside the customer's Azure tenant, with no public model API and no full-codebase indexing by default.",
            "Evidence first: the output is not just a suggestion stream. It is a signed, replayable record of what was generated, tested, rejected, and proposed.",
            "Gated quality: candidate tests must compile, run, pass, stay non-flaky, increase changed-line coverage, and catch mutation-injected faults before they are shown to humans.",
            "Human control: Sentinel never auto-merges code. The safest message is that AI accelerates preparation, while engineers retain authority.",
            "Multi-platform ambition: the PRD covers backend, web, iOS, Android, React Native/Flutter, E2E, and visual regression from one orchestrator.",
            "Regulated delivery fit: the architecture speaks the language of IEC 62304, FDA 21 CFR Part 11, and FDA Computer Software Assurance: risk-based evidence, audit trails, and controlled change.",
        ],
    )

    add_heading(doc, "The Authentic Leadership Narrative", 1)
    narrative = [
        ("What we should say", "We are not trying to replace engineers or reviewers. We are trying to make every review start with better test evidence."),
        ("What we should avoid", "Do not claim Sentinel eliminates defects, replaces QA, or makes regulatory work automatic. Those claims will sound inflated."),
        ("Why now", "AI is increasing code volume. The winning engineering organizations will add AI quality controls, not only AI code generation."),
        ("Why Stryker", "Stryker already values innovation, product quality, and patient impact. Sentinel supports that culture by adding disciplined speed."),
    ]
    add_table(doc, ["Point", "Leadership-ready wording"], narrative, [1.4, 4.6], 9.2)

    add_heading(doc, "Risks To Name Up Front", 1)
    risks = [
        (
            "Security approval",
            "Azure AI Foundry and Claude availability must be confirmed with Stryker security and procurement.",
            "Make this Discovery Call 1. Offer approved-model fallback if needed.",
        ),
        (
            "Generated-test noise",
            "Bad AI tests can slow teams down if they reach reviewers.",
            "Keep the six gates strict. Dropped tests should be invisible by default.",
        ),
        (
            "iOS CI constraints",
            "iOS UI testing needs macOS runners and Xcode, which may not be available in all CI environments.",
            "Pilot backend first, then add iOS after runner path is confirmed.",
        ),
        (
            "Validation burden",
            "If used in regulated workflows, Sentinel itself may need tool validation or CSA evidence.",
            "Treat validation as a feature, not a surprise. Include IQ/OQ/PQ or CSA package planning.",
        ),
        (
            "ROI proof",
            "Leadership will want measured improvement, not enthusiasm.",
            "Use a 90-day pilot with clear before/after metrics.",
        ),
    ]
    add_table(doc, ["Risk", "Honest concern", "Practical answer"], risks, [1.2, 2.45, 2.35], 8.5)

    add_heading(doc, "90-Day Pilot Recommendation", 1, page_break_before=True)
    add_body(
        doc,
        "Start with a controlled pilot that is useful even if Sentinel is not rolled out immediately. The pilot should produce a decision memo, not just a demo.",
    )
    pilot = [
        ("Weeks 1-2", "Discovery", "Security, platform, engineering, and quality alignment. Confirm repos, CI, model approval, audit needs, and baseline metrics."),
        ("Weeks 3-6", "Backend PR pilot", "Run Sentinel on a backend service for unit and integration tests. Keep all output in patch PRs for human review."),
        ("Weeks 7-10", "Web or API expansion", "Add component/E2E coverage and failure reporting. Measure false positives and reviewer acceptance."),
        ("Weeks 11-12", "Leadership readout", "Present metrics, accepted tests, rejected tests, evidence samples, risks, and a scale/no-scale recommendation."),
    ]
    add_table(doc, ["Timing", "Phase", "What happens"], pilot, [1.0, 1.35, 3.65], 8.8)

    add_heading(doc, "Metrics Leadership Will Trust", 1)
    add_bullets(
        doc,
        [
            "Changed-line test coverage before and after Sentinel.",
            "Time from PR opened to Sentinel verdict.",
            "Generated-test acceptance rate after human review.",
            "Flake rate of accepted generated tests.",
            "Number of mutation-injected faults caught by generated tests.",
            "Review-to-merge latency compared with baseline.",
            "Quality and compliance feedback from engineering, platform, security, and quality stakeholders.",
        ],
    )

    add_heading(doc, "Suggested Meeting Talk Track", 1)
    add_numbered(
        doc,
        [
            "Open with the real problem: AI is helping teams create code faster, but review evidence is still manual and uneven.",
            "Acknowledge the market: there are good tools for code review, test generation, visual testing, and QA automation.",
            "Explain the wedge: Sentinel is a private PR quality gate that creates only verified test evidence and keeps humans in control.",
            "Tie to Stryker: SGTC can turn this into a repeatable capability for regulated software teams across divisions.",
            "Ask for a pilot: one focused repo, clear success metrics, security and quality involved from day one.",
        ],
    )

    add_heading(doc, "Simple Closing Ask", 1)
    add_callout(
        doc,
        "Ask",
        "Approve a 90-day Sentinel pilot with one engineering team, one platform owner, one security reviewer, and one quality/compliance reviewer. The output will be a measured recommendation, not a sales demo.",
        "FFFBEB",
    )

    add_heading(doc, "Research Sources", 1)
    add_body(
        doc,
        "Sources reviewed for the competitor and market summary. These links are included so the pitch can be defended in follow-up conversations.",
    )
    sources = [
        ("Stryker SGTC R&D", "https://www.stryker.com/in/en/about/our-locations/stryker-global-technology-center-r-and-d.html"),
        ("Stryker SmartHospital Platform, Mar. 2026", "https://www.stryker.com/us/en/smart-care/news/stryker-launches-smarthospital-platform.html"),
        ("Diffblue Cover docs", "https://docs.diffblue.com/get-started/what-is-diffblue-cover"),
        ("Qodo generate tests docs", "https://docs.qodo.ai/qodo-documentation/qodo-merge/tools/tools-list/generate-tests/"),
        ("CodeRabbit unit-test generation docs", "https://docs.coderabbit.ai/finishing-touches/unit-test-generation/"),
        ("GitHub Copilot coding agent docs", "https://docs.github.com/en/copilot/concepts/coding-agent/about-copilot-coding-agent"),
        ("mabl AI test automation", "https://www.mabl.com/ai-test-automation"),
        ("Tricentis Tosca", "https://www.tricentis.com/products/automate-continuous-testing-tosca"),
        ("Launchable predictive test selection", "https://www.launchableinc.com/docs/features/predictive-test-selection/"),
        ("Applitools Eyes", "https://applitools.com/platform/eyes/"),
        ("FDA Computer Software Assurance guidance", "https://www.fda.gov/regulatory-information/search-fda-guidance-documents/computer-software-assurance-production-and-quality-management-system-software"),
        ("FDA Part 11 guidance", "https://www.fda.gov/regulatory-information/search-fda-guidance-documents/part-11-electronic-records-electronic-signatures-scope-and-application"),
        ("FDA IEC 62304 recognition", "https://www.accessdata.fda.gov/scripts/cdrh/cfdocs/cfstandards/detail.cfm?standard__identification_no=38829"),
    ]
    add_table(doc, ["Source", "URL"], sources, [1.75, 4.25], 7.7)

    doc.core_properties.title = "Sentinel Stryker Leadership Pitch"
    doc.core_properties.subject = "Research-backed executive memo"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
